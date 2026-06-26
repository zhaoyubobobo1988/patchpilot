"""生成架构盘点报告 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.0)

# ── 样式工具 ──────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None):
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rPr.insert(0, rFonts)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 73, 125))
    # 下边框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(68, 114, 196))
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(0, 0, 0))
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # 浅灰底色段落
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 表头
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=(255,255,255))
        # 蓝色背景
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
    # 数据行
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9E2F3')
                tcPr.append(shd)
    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(80)
r = cover.add_run('OpenClaw 架构盘点报告')
set_font(r, size=24, bold=True, color=(31,73,125))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(12)
r2 = sub.add_run('多 Agent 系统成熟度评估与改造路线图')
set_font(r2, size=14, color=(89,89,89))

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_before = Pt(24)
r3 = date_p.add_run(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
set_font(r3, size=11, color=(89,89,89))

note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = note_p.add_run('基于真实代码阅读 · 不含任何推测 · 未修改任何文件')
set_font(r4, size=10, color=(128,128,128))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. 执行摘要
# ══════════════════════════════════════════════════════════════
h1('1. 执行摘要')
body('当前系统本质上是一条 9 阶段固定顺序的 DAG 工作流，所有控制流写死在 pipeline.py 的 run_pipeline() 函数中。每个被称为"Agent"的组件实质上是一次 LLM 调用的包装类（类型 A），没有任何组件拥有自主决策权。Agent 之间完全不互相通信，所有数据通过 pipeline.py 中的 Python 变量显式传递。状态仅存在于单次进程内存中，进程退出即丢失，无法恢复。')
body('')
body('Multi-Agent 成熟度整体评分：1.5 / 5（详见第 16 节）', indent=1)

# ══════════════════════════════════════════════════════════════
# 2. 核心目录树
# ══════════════════════════════════════════════════════════════
h1('2. 核心目录树')
code_block(
    'myownagent/\n'
    '├── pipeline.py                    # 唯一程序入口，控制流全在这里\n'
    '│\n'
    '├── agents/                        # Agent 实现层\n'
    '│   ├── base.py                    # AgentTask / AgentResult / AgentAdapter Protocol\n'
    '│   ├── claude_code.py             # ClaudeCodeAgent（claude CLI subprocess 封装）\n'
    '│   ├── codex.py                   # CodexAgent（codex CLI subprocess 封装，备用）\n'
    '│   ├── registry.py                # AgentRegistry（name→instance 查找表）\n'
    '│   ├── router.py                  # AgentRouter（仅用于 Reviewer backend 选择）\n'
    '│   ├── context_agent/agent.py     # ContextAgent（纯 Python，无 LLM）\n'
    '│   ├── orchestrator/agent.py      # OrchestratorAgent（一次性计划生成）\n'
    '│   ├── test_agent/agent.py        # TestAgent（litellm，生成 pytest 文件）\n'
    '│   ├── worker/agent.py            # ClaudeCodeWorker（claude CLI subprocess）\n'
    '│   ├── worker/workspace.py        # WorkerWorkspaceManager（clone/worktree 隔离）\n'
    '│   ├── aggregator/agent.py        # AggregatorAgent（litellm，合并 patch）\n'
    '│   ├── integrator/agent.py        # IntegratorAgent（纯 Python 规则校验，无 LLM）\n'
    '│   ├── review_agent/agent.py      # ReviewAgent（claude CLI subprocess）\n'
    '│   ├── github_agent/agent.py      # GitHubAgent（git subprocess + GitHub REST API）\n'
    '│   └── debug_agent/agent.py       # DebugAgent（litellm，CI 失败修复）\n'
    '│\n'
    '├── models/                        # 数据结构层（Pydantic）\n'
    '│   ├── task.py                    # SubTask / FeatureTask / TaskGraph\n'
    '│   ├── patch.py                   # PatchResult / MergedPatch / ReviewResult\n'
    '│   ├── context.py                 # AgentContext / CodeContext / PipelineRun\n'
    '│   └── github.py                  # PRRequest / PRResult / CICheckResult\n'
    '│\n'
    '├── config/\n'
    '│   ├── settings.py                # Pydantic Settings（读 .env）\n'
    '│   └── llm.py                     # llm_complete()（litellm 统一入口）\n'
    '│\n'
    '├── telemetry/\n'
    '│   ├── execution_log.py           # JSONL append-only 写入\n'
    '│   └── log_stats.py               # 跨 run 统计分析\n'
    '│\n'
    '└── tests/\n'
    '    ├── unit/                      # 单元测试（~18 个文件）\n'
    '    └── integration/test_smoke.py  # E2E smoke test'
)

h2('入口说明')
add_table(
    ['项目', '位置'],
    [
        ['程序入口', 'pipeline.py::main() → asyncio.run(run_pipeline(...))'],
        ['需求进入点', 'pipeline.py:170  run_pipeline(raw_requirement, repository, ...)'],
        ['PR 创建点', 'agents/github_agent/agent.py:74  GitHubAgent.create_pr()'],
        ['状态容器', 'models/context.py  PipelineRun（进程内存）'],
    ],
    col_widths=[3.5, 13]
)

# ══════════════════════════════════════════════════════════════
# 3. 真实架构图
# ══════════════════════════════════════════════════════════════
h1('3. 当前真实架构图')
code_block(
    '【单进程 Python asyncio】\n\n'
    'pipeline.py::run_pipeline()   ← 持有所有控制权\n'
    '│\n'
    '├─ 持有：run_id, model, task, ctx, run（PipelineRun）\n'
    '│\n'
    '├─ 所有 Agent 都是 pipeline.py 直接 new 出来并调用的\n'
    '│\n'
    '└─ 数据流：pipeline.py 变量 → 参数传入 → 返回值 → 下一步\n\n'
    'Agent 之间：零通信，无消息，无事件，无共享状态\n\n'
    '谁决定下一步？   pipeline.py（100%）\n'
    '谁管理状态？     pipeline.py 中的局部变量\n'
    '谁负责重试？     pipeline.py 中的 for 循环\n'
    '谁负责终止？     pipeline.py 中的 return / raise'
)

# ══════════════════════════════════════════════════════════════
# 4. 完整执行链路
# ══════════════════════════════════════════════════════════════
h1('4. 当前完整执行链路')
code_block(
    'Requirement (CLI 参数)\n'
    '  │\n'
    '  ▼ [Stage 0 - pipeline.py:185]\n'
    'git clone 目标仓库 → C:/tmp/openclaw-workspaces/{run_id}/\n'
    '  │\n'
    '  ▼ [Preflight]\n'
    'WorkerWorkspaceManager.validate_strategy()  ← 纯校验，失败即终止\n'
    '  │\n'
    '  ▼ [Stage 1a - pipeline.py:211]\n'
    'ContextAgent.gather()                        ← 纯 Python，无 LLM\n'
    '  │   输出：CodeContext（最多 20 个相关文件）\n'
    '  ▼\n'
    '[Stage 1b - pipeline.py:221]\n'
    'OrchestratorAgent.decompose()               ← claude CLI（→ litellm fallback）\n'
    '  │   输出：TaskGraph（subtasks + parallel_groups）\n'
    '  ▼\n'
    '[Stage 2 - pipeline.py:232]\n'
    'asyncio.gather(*[TestAgent.generate(st) for st in all_subtasks])  ← litellm 并行\n'
    '  │   每个 subtask 写入 features/{feature}/test_{id}.py\n'
    '  ▼\n'
    '[Stage 3 - pipeline.py:248]\n'
    'for group in task_graph.parallel_groups:\n'
    '    asyncio.gather(*[worker.execute(st) for ...])  ← claude CLI 并行\n'
    '  │   每 Worker 独立 clone workspace，输出 PatchResult（git diff）\n'
    '  ▼\n'
    '[Stage 4a] AggregatorAgent.merge()          ← litellm（冲突时）\n'
    '  ▼\n'
    '[Stage 4b] IntegratorAgent.integrate()      ← 纯 Python，4 项规则校验\n'
    '  │   失败 → ValueError → pipeline 终止\n'
    '  ▼\n'
    '[Stage 5] ReviewAgent.review()              ← claude CLI（最多 3 次）\n'
    '  ├─ approved → 进入 Stage 6\n'
    '  └─ blocked  → 重跑所有 Worker → Aggregator → Integrator → Review\n'
    '                超过 2 次重试 → pipeline 终止\n'
    '  ▼\n'
    '[Stage 6] GitHubAgent.apply_and_push()      ← git subprocess\n'
    '          GitHubAgent.create_pr()            ← GitHub REST API\n'
    '  ▼\n'
    '[Stage 7] CI 轮询 + DebugAgent 重试（最多 6 次）\n'
    '  ├─ CI SUCCESS → run.ci_passed=True\n'
    '  └─ CI FAILURE → DebugAgent.fix() → 新 PR → 继续轮询\n'
    '  ▼\n'
    'run.stage = "done"\n'
    'return PipelineRun'
)

h2('串行 / 并行 / 重试一览')
add_table(
    ['阶段', '执行方式', '并行', '重试', '失败处理'],
    [
        ['Stage 0: clone', 'subprocess', '否', '否', '崩溃终止'],
        ['Stage 1a: context', 'Python', '否', '否', '返回空结果'],
        ['Stage 1b: orchestrate', 'claude CLI', '否', '否（有 litellm fallback）', '崩溃终止'],
        ['Stage 2: test-gen', 'litellm', '是（asyncio）', '否', '写空文件'],
        ['Stage 3: worker', 'claude CLI', '是（同组内）', '否', 'PatchResult.FAILED'],
        ['Stage 4a: aggregate', 'litellm', '否', '否（最长 hunk fallback）', 'MergedPatch.FAILED'],
        ['Stage 4b: integrate', 'Python 规则', '否', '否', '抛 ValueError 终止'],
        ['Stage 5: review', 'claude CLI', '否', '是（最多 2 次）', '终止 pipeline'],
        ['Stage 6: github', 'git + HTTP', '否', '否', '崩溃终止'],
        ['Stage 7: CI + debug', 'litellm + HTTP 轮询', '否', '是（最多 5 次）', '标记 ci_passed=False'],
    ],
    col_widths=[3.5, 3, 1.5, 2, 5.5]
)

# ══════════════════════════════════════════════════════════════
# 5. Agent 能力对比表
# ══════════════════════════════════════════════════════════════
h1('5. Agent 能力对比表')
add_table(
    ['Agent', '调用 LLM', '调用工具', '多轮运行', '保存状态', '能委派任务', '类型'],
    [
        ['ContextAgent', '否', '否（Python AST）', '否', '否', '否', '纯函数'],
        ['OrchestratorAgent', '是（claude CLI）', '是（CLI 内部）', '否（单次）', '否', '否', 'A'],
        ['TestAgent', '是（litellm）', '否', '否', '否', '否', 'A'],
        ['ClaudeCodeWorker', '是（claude CLI）', '是（CLI 内部）', '否（单次）', '否', '否', 'A/B 混合'],
        ['AggregatorAgent', '是（litellm）', '否', '否', '否', '否', 'A'],
        ['IntegratorAgent', '否', '是（可选 shell）', '否', '是（last_result）', '否', '纯规则'],
        ['ReviewAgent', '是（claude CLI）', '否', '否', '是（last_route）', '否', 'A'],
        ['GitHubAgent', '否', '是（git+HTTP）', '是（CI 轮询）', '否', '否', '工具执行器'],
        ['DebugAgent', '是（litellm）', '否', '否', '否', '否', 'A'],
    ],
    col_widths=[3.5, 2.5, 2.5, 2, 2, 2.5, 2.5]
)
body('结论：绝大多数 Agent 属于类型 A（单次 LLM 调用包装器）。Worker 内部的 claude CLI 进程有工具调用能力，但从 Python 层看，Worker 仍是"调用一次，等结果"。')

# ══════════════════════════════════════════════════════════════
# 6. Orchestrator 深度分析
# ══════════════════════════════════════════════════════════════
h1('6. Orchestrator 深度分析')
h2('6.1 计划时机')
body('仅在 pipeline.py:223 调用一次 orchestrator.decompose(task)，执行过程中不再调用。Orchestrator 不知道 Worker 是否成功、Review 是否通过、CI 是否失败。')

h2('6.2 输出示例（真实运行）')
code_block(
    '{\n'
    '  "feature_name": "auth-login-logging",\n'
    '  "subtasks": [\n'
    '    {\n'
    '      "id": "task-01",\n'
    '      "feature": "auth",\n'
    '      "goal": "在文件顶部 import logging",\n'
    '      "files": ["features/auth/login.py"],\n'
    '      "constraints": []\n'
    '    },\n'
    '    {\n'
    '      "id": "task-02",\n'
    '      "feature": "auth",\n'
    '      "goal": "在函数内加 logging.info(\'login attempt: %s\', username)",\n'
    '      "files": ["features/auth/login.py"],\n'
    '      "constraints": []\n'
    '    }\n'
    '  ],\n'
    '  "parallel_groups": [["task-01", "task-02"]],\n'
    '  "dependencies": {}\n'
    '}'
)

h2('6.3 TaskGraph 真实字段')
body('feature_task.id / raw_requirement / feature_name / repository / base_branch')
body('feature_task.subtasks[].id / feature / goal / files / constraints')
body('feature_task.subtasks[].status  ← 固定 PENDING，从未更新（重要 Bug）')
body('feature_task.subtasks[].assigned_worker_id  ← 固定 None，从未赋值（重要 Bug）')
body('parallel_groups / dependencies')

h2('6.4 Orchestrator 当前能力')
add_table(
    ['能力', '是否存在', '说明'],
    [
        ['动态新增任务', '否', '计划只生成一次'],
        ['删除 / 修改任务', '否', '同上'],
        ['选择不同 Agent 执行', '否', 'Worker 固定为 ClaudeCodeWorker'],
        ['根据 Worker 发现重规划', '否', 'Worker 结果不回传 Orchestrator'],
        ['根据测试失败重规划', '否', 'Orchestrator 不感知测试结果'],
        ['请求 ContextAgent 补充', '否', '无此调用路径'],
        ['请求人工介入', '否', '不存在'],
        ['输出解析失败处理', '否', 'json.loads 失败直接 crash，无兜底计划'],
    ],
    col_widths=[5, 2.5, 9]
)

# ══════════════════════════════════════════════════════════════
# 7. Worker 与并发分析
# ══════════════════════════════════════════════════════════════
h1('7. Worker 与并发分析')
add_table(
    ['项目', '状态', '说明'],
    [
        ['Worker 数量决策', 'Orchestrator 的 parallel_groups', 'pipeline.py 不做额外限制'],
        ['MAX_PARALLEL_WORKERS', '配置存在但未使用', 'settings.py 有该字段，pipeline.py 未引用'],
        ['并发机制', 'asyncio.gather + 子进程', '协程并发，每 Worker 独立 claude 进程'],
        ['Workspace 隔离', '是', '{run_id}-worker-{i}/ 独立 clone/worktree'],
        ['Worker 输出格式', 'unified diff 文本', 'git diff --cached 输出'],
        ['Worker 上报新任务', '否', '只能返回 PatchResult(status=FAILED)'],
        ['同文件并发修改', '存在风险', 'Aggregator 做 LLM 合并，但 _collect_worker_changes 用 copy 覆盖'],
    ],
    col_widths=[4, 4, 8.5]
)

h2('潜在并发问题')
add_table(
    ['问题', '是否存在', '影响'],
    [
        ['文件冲突', '存在', '同组 Worker 改同一文件，Aggregator 做 LLM 合并'],
        ['_collect_worker_changes 覆盖 Bug', '存在（严重）', '后扫描到的 Worker workspace 文件覆盖前者，一个 Worker 改动丢失'],
        ['依赖执行错误', '存在', 'dependencies 字段有值但未校验，有依赖的 subtask 仍可能并行'],
        ['上下文覆盖', '不存在', '独立 workspace 隔离'],
    ],
    col_widths=[5, 2.5, 9]
)

# ══════════════════════════════════════════════════════════════
# 8. 状态管理分析
# ══════════════════════════════════════════════════════════════
h1('8. 状态管理分析')
add_table(
    ['状态对象', '存储位置', '内容', '持久化'],
    [
        ['PipelineRun', 'Python 内存', 'run_id, stage, pr_url, ci_passed, error_log', '否'],
        ['AgentContext', 'Python 内存', 'run_id, repository, workspace_path, model', '否'],
        ['FeatureTask + TaskGraph', 'Python 内存', '原始需求, subtasks, parallel_groups', '否'],
        ['MergedPatch', 'Python 内存', 'merged_diff（随 stage 覆盖）', '否'],
        ['JSONL log (exec.jsonl)', '文件系统', 'ExecutionRecord 事件流（无输入输出内容）', '是'],
    ],
    col_widths=[3.5, 3, 8, 2]
)

h2('可恢复性')
add_table(
    ['能力', '是否存在'],
    [
        ['服务重启后任务恢复', '否'],
        ['支持暂停和继续', '否'],
        ['从失败节点继续', '否'],
        ['Checkpoint', '否'],
        ['统一事件模型', '部分（JSONL，仅后验分析）'],
    ],
    col_widths=[8, 8.5]
)

# ══════════════════════════════════════════════════════════════
# 9. Agent 间通信分析
# ══════════════════════════════════════════════════════════════
h1('9. Agent 间通信分析')
add_table(
    ['通信机制', '是否存在'],
    [
        ['直接 Agent-to-Agent 消息', '否'],
        ['共享消息历史', '否'],
        ['共享黑板 Blackboard', '否'],
        ['事件总线', '否'],
        ['Pub/Sub', '否'],
        ['任务委派', '否'],
        ['Handoff', '否'],
        ['Agent 主动调用另一个 Agent', '否'],
        ['Agent 发布新发现', '否'],
        ['Agent 请求其他 Agent 帮助', '否'],
    ],
    col_widths=[8, 8.5]
)
body('当前实际通信模式：pipeline.py 调用 AgentA.method(input) → 拿到 result_A → 调用 AgentB.method(result_A)')
body('这是"A 的返回值被 Pipeline 传给 B"，不是"A 主动调用 B"。')
body('')
body('系统分类判断：类型 1 —— 多个 LLM 节点的固定 Workflow')
body('判断理由：控制流 100% 在 pipeline.py 代码中，无动态路由；没有任何 Agent 持有目标、能自主规划下一步；Agent 间零通信；AgentRouter 仅用于 Reviewer backend 选择，不是任务路由。')

# ══════════════════════════════════════════════════════════════
# 10. 模型与 Claude CLI 调用分析
# ══════════════════════════════════════════════════════════════
h1('10. 模型与 Claude CLI 调用分析')
h2('10.1 两条调用路径')
add_table(
    ['路径', '使用角色', '命令/方式', 'Prompt 传入', '输出解析'],
    [
        ['claude CLI subprocess', 'Orchestrator(主), Worker(主), Reviewer(主)', 'claude --print --output-format json|text --dangerously-skip-permissions --bare', 'stdin', 'JSON 信封解析或 raw stdout'],
        ['litellm', 'Orchestrator(fallback), TestAgent, Aggregator, Reviewer(fallback), Debug', 'litellm.acompletion()', 'messages 参数', '直接取 choices[0].message.content'],
    ],
    col_widths=[2.5, 4, 4.5, 2.5, 3]
)

h2('10.2 关键能力项')
add_table(
    ['项目', '状态'],
    [
        ['JSON Schema / Pydantic 校验', '否（json.loads 后无 schema 校验）'],
        ['每 Agent 独立 CLI 进程', '是（无会话连续性，--bare 标志）'],
        ['不同 Agent 使用不同模型', '否（共用 LLM_MODEL / CLAUDE_CODE_MODEL）'],
        ['超时', '是（subprocess: CLAUDE_CODE_TIMEOUT=300s；litellm: 无）'],
        ['重试', '否（单次调用，失败即返回）'],
        ['Token / 成本统计', '否'],
        ['模型调用与业务逻辑解耦', '部分（AgentAdapter Protocol 存在，但 litellm 未实现该 Protocol）'],
    ],
    col_widths=[6, 10.5]
)

# ══════════════════════════════════════════════════════════════
# 11. 测试与失败恢复分析
# ══════════════════════════════════════════════════════════════
h1('11. 测试与失败恢复分析')
add_table(
    ['项目', '状态', '说明'],
    [
        ['TestAgent 执行测试', '否', '只生成 pytest 文件，不运行'],
        ['Worker 完成后自动运行测试', '否', '无此机制'],
        ['IntegratorAgent 运行测试', '可选（默认跳过）', 'INTEGRATION_TEST_COMMAND 默认为空'],
        ['CI 测试', '是', 'GitHubAgent.poll_ci() 轮询 GitHub check-runs API'],
        ['DebugAgent 失败类型区分', '否', '全部 CI 日志一并发给 LLM，不做分类'],
        ['Review 失败后退回指定 Worker', '否', '重跑所有 Worker（粒度过粗）'],
        ['最大 Review 重试次数', '2 次', '_MAX_REVIEW_RETRIES = 2（pipeline.py:41）'],
        ['最大 Debug 重试次数', '5 次', 'MAX_DEBUG_RETRIES = 5（settings.py）'],
    ],
    col_widths=[4.5, 3, 9]
)

# ══════════════════════════════════════════════════════════════
# 12. Git / GitHub 分析
# ══════════════════════════════════════════════════════════════
h1('12. Git / GitHub 分析')
add_table(
    ['操作', '执行者', '代码位置'],
    [
        ['git clone（主 workspace）', 'pipeline.py _clone_repo()', 'pipeline.py:66'],
        ['git clone --local（Worker 隔离）', 'WorkerWorkspaceManager', 'workspace.py:193'],
        ['git checkout -b（feature 分支）', 'GitHubAgent', 'github_agent/agent.py:36'],
        ['收集 Worker 改动（文件 copy）', 'GitHubAgent._collect_worker_changes()', 'github_agent/agent.py:163'],
        ['git add -A + commit + push', 'GitHubAgent', 'github_agent/agent.py:52-63'],
        ['创建 Draft PR', 'GitHubAgent.create_pr()', 'github_agent/agent.py:74'],
        ['CI 轮询', 'GitHubAgent.poll_ci()', 'github_agent/agent.py:99'],
    ],
    col_widths=[5, 4.5, 7]
)

h2('高风险操作检查')
add_table(
    ['操作', '是否存在'],
    [
        ['git reset --hard', '否'],
        ['git push --force', '否'],
        ['git clean -fd', '否'],
        ['人工审批节点', '否（PR 为 Draft，但无强制 review 步骤）'],
    ],
    col_widths=[8, 8.5]
)

body('重要 Bug：_collect_worker_changes 使用 shutil.copy2 文件覆盖，当两个 Worker 都修改了同一文件时，后扫描到的 Worker 文件会完全覆盖前者，导致一个 Worker 的改动丢失。')

# ══════════════════════════════════════════════════════════════
# 13. 权限与安全分析
# ══════════════════════════════════════════════════════════════
h1('13. 权限与安全分析')
add_table(
    ['工具', '可读文件', '可写文件', '可执行命令', '访问网络', '操作 Git'],
    [
        ['ClaudeCodeAgent (Worker)', '是（workspace）', '是（workspace）', '是（无限制）', '是', '是'],
        ['ClaudeCodeAgent (Orchestrator/Reviewer)', '是（workspace）', '受角色限制', '受角色限制', '是', '否'],
        ['GitHubAgent', '是', '是', '是（git）', '是（GitHub API）', '是（push 权限）'],
        ['IntegratorAgent', '否', '否', '是（可选 shell）', '否', '否'],
    ],
    col_widths=[4.5, 2, 2, 2.5, 2, 3.5]
)

h2('安全问题')
bullet('Worker 使用 --dangerously-skip-permissions：claude CLI 进程在 workspace 内无任何工具限制')
bullet('文件路径限制：仅有 Python 层 _validate_diff() 软校验（features/ 前缀），claude 进程本身不受路径约束')
bullet('无沙箱/容器隔离：Worker 子进程可访问宿主机网络和文件系统')
bullet('GitHubAgent 无人工审批：直接 push 并创建 PR，无强制拦截')

# ══════════════════════════════════════════════════════════════
# 14. 可观测性与成本分析
# ══════════════════════════════════════════════════════════════
h1('14. 可观测性与成本分析')
add_table(
    ['记录项', '是否存在', '位置'],
    [
        ['统一 run_id', '是', 'ExecutionRecord.run_id'],
        ['Agent 执行时间', '是（elapsed_seconds）', 'execution_log.py'],
        ['Agent 输入（prompt）', '否', '故意排除'],
        ['Agent 输出', '否', '故意排除'],
        ['Token 数', '否', '未接入'],
        ['API 成本', '否', '未接入'],
        ['重试次数', '是', 'pipeline_completed.metadata.debug_retry_count'],
        ['PR 地址', '是', 'pipeline_completed.metadata.pr_url'],
        ['错误信息', '是（截断 1000 字符）', 'ExecutionRecord.error'],
        ['Git diff 内容', '否', '未记录'],
    ],
    col_widths=[5, 2.5, 9]
)
body('无法统计一次完整 Run 的 Token 消耗和 API 成本。无法回放历史执行（JSONL 不含输入输出内容）。MAX_PARALLEL_WORKERS 配置有值但未生效。')

# ══════════════════════════════════════════════════════════════
# 15. 当前十大架构问题
# ══════════════════════════════════════════════════════════════
h1('15. 当前十大架构问题')

problems = [
    ('控制流完全硬编码',
     'pipeline.py::run_pipeline() 是线性函数，9 个 stage 写死，无任何动态路由',
     '无法适应需求变化，任何新场景必须改 pipeline.py',
     '高',
     '提取 Orchestrator Loop，让 Orchestrator 在每个 stage 后决策下一步'),
    ('Orchestrator 一次性规划，无法响应执行结果',
     'orchestrator.decompose() 只调用一次（pipeline.py:223），此后不再参与',
     'Worker 失败、测试失败、新发现均无法触发重规划',
     '高',
     '实现 Orchestrator 反馈循环，接收 Worker/Reviewer 上报事件'),
    ('SubTask 状态字段存在但从未更新',
     'SubTask.status 固定 PENDING，assigned_worker_id 固定 None',
     '系统无法跟踪真实任务状态',
     '中',
     'Worker 启动/完成时更新 subtask.status，建立真实 TaskBoard'),
    ('Worker 发现新问题无法上报',
     'ClaudeCodeWorker.execute() 只能返回 PatchResult(status=FAILED)',
     'Worker 发现依赖缺失等情况只能静默失败',
     '高',
     '增加 AgentEvent 上报机制，Worker 可以发布 NewTaskDiscovered 事件'),
    ('状态全在内存，进程重启即丢失',
     'run, ctx, task_graph 均为 pipeline.py 局部变量，无持久化',
     '长时间任务崩溃后必须从头开始',
     '高',
     '建立 RunState 持久化（SQLite 或文件），支持 checkpoint 恢复'),
    ('_collect_worker_changes 文件 copy 覆盖导致改动丢失',
     'github_agent/agent.py:163 使用 shutil.copy2，同一文件后者覆盖前者',
     '两个 Worker 改同一文件时，一个 Worker 改动丢失',
     '高',
     '改用 git merge 或应用 AggregatorAgent 生成的 merged_diff'),
    ('无 Token / 成本跟踪',
     'execution_log.py 有意排除 token 计数，config/llm.py 不返回 usage',
     '无法评估和控制每次 Run 的 API 成本',
     '中',
     '从 litellm response.usage 提取 token 数，写入 JSONL'),
    ('MAX_PARALLEL_WORKERS 配置未使用',
     'settings.py 有 MAX_PARALLEL_WORKERS=4，pipeline.py 未引用',
     '如 Orchestrator 生成 10 个并行子任务，同时启动 10 个 claude 进程',
     '中',
     '在 _run_workers 增加 semaphore 限制'),
    ('Review 失败只能重跑所有 Worker，粒度过粗',
     'pipeline.py:295 对所有 subtask 添加相同约束，重跑所有 Worker',
     '只有 1 个文件有问题时仍重新生成所有 patch，浪费资源',
     '中',
     'Review 反馈精确到 subtask 级别，只重跑问题 subtask'),
    ('claude CLI 调用无重试',
     'agents/claude_code.py 无重试逻辑，失败直接走 litellm fallback',
     '网络抖动导致直接降级，可靠性低',
     '低',
     '在 ClaudeCodeAgent 内部增加 1-2 次指数退避重试'),
]

for i, (title, evidence, impact, severity, suggestion) in enumerate(problems, 1):
    h2(f'问题 {i}：{title}')
    add_table(
        ['项目', '内容'],
        [
            ['证据', evidence],
            ['影响', impact],
            ['严重程度', severity],
            ['建议方向', suggestion],
        ],
        col_widths=[2.5, 14]
    )

# ══════════════════════════════════════════════════════════════
# 16. Multi-Agent 成熟度评分
# ══════════════════════════════════════════════════════════════
h1('16. Multi-Agent 成熟度评分')
add_table(
    ['能力', '分数', '判断依据'],
    [
        ['Agent 自主决策', '0/5', '无任何 Agent 能决定下一步；控制流 100% 在 pipeline.py'],
        ['动态路由', '1/5', 'AgentRouter 仅用于 Reviewer backend 选择，不是任务路由'],
        ['动态任务拆分', '0/5', 'Orchestrator 一次性规划，执行中不追加任务'],
        ['Agent 间通信', '0/5', '零 Agent-to-Agent 通信，全部通过 pipeline.py 变量传递'],
        ['新发现上报', '0/5', '无上报机制，Worker 只能 FAILED 或 SUCCESS'],
        ['并发执行', '3/5', 'asyncio.gather + 独立 workspace，有真实并发；受 parallel_groups 固定约束'],
        ['失败自主恢复', '1/5', '有重试循环但固定（Review×2, Debug×5），无分类恢复'],
        ['状态持久化', '1/5', '仅有 JSONL 日志；无 checkpoint，无恢复能力'],
        ['可观测性', '2/5', '有 run_id + JSONL 事件流 + 统计分析；缺 token/cost/输入输出'],
        ['成本控制', '0/5', '无 Token 统计，MAX_PARALLEL_WORKERS 未使用'],
        ['安全边界', '1/5', '仅有 features/ 路径软校验；Worker 不受沙箱约束'],
        ['生产可恢复性', '0/5', '进程崩溃后无法恢复，无 checkpoint'],
    ],
    col_widths=[3.5, 1.5, 11.5]
)

body('')
body('总分：9/60 = 1.5/5')
body('当前项目更接近：固定 Workflow（类型 1），即多个 LLM 节点的固定 DAG')

# ══════════════════════════════════════════════════════════════
# 17. 最小改造路线图
# ══════════════════════════════════════════════════════════════
h1('17. 最小改造路线图')
body('原则：保留现有 Agent 类，优先解决控制流和状态，不引入外部框架，每 PR 可独立验收。')

prs = [
    ('PR 1', '统一 Agent 输出协议 + 增加 AgentEvent',
     'agents/base.py（AgentResult 增加 events 字段）\n新增 models/events.py（AgentEvent 数据结构）',
     '所有单元测试通过；AgentResult 向后兼容',
     'events 默认为空列表，现有调用方不受影响'),
    ('PR 2', '建立统一 RunState 并持久化',
     'models/context.py（PipelineRun 增加 task_graph_snapshot / stage_timestamps）\n新增 state/run_state.py（RunStateManager）\npipeline.py（各 stage 结束后调用 save）',
     '任务完成后 {workspace}/run_state.json 文件存在且完整',
     'save 失败不影响主流程（非阻塞）'),
    ('PR 3', '修复 _collect_worker_changes 文件覆盖 Bug',
     'agents/github_agent/agent.py::_collect_worker_changes()',
     '两个 Worker 改同一文件，最终 commit 包含两者改动',
     '需要理解 AggregatorAgent 的 merged_diff 与文件 apply 的关系'),
    ('PR 4', '补全 SubTask 状态更新',
     'pipeline.py::_run_workers()\nagents/worker/agent.py',
     'task_graph.feature_task.subtasks[*].status 反映真实执行结果',
     '无'),
    ('PR 5', '将 pipeline.py 的 Stage 循环提取为 StageExecutor',
     '新增 pipeline/executor.py（StageExecutor 类）\npipeline.py（变为调度者）',
     '功能不变，各 Stage 可单独测试',
     '较大重构，需完整回归测试'),
    ('PR 6', 'Orchestrator 接收 AgentEvent，支持追加任务',
     'agents/orchestrator/agent.py（增加 handle_events() 方法）\npipeline.py（Worker 完成后传递 events）',
     'Worker 上报 NEW_SUBTASK_DISCOVERED 事件时能执行新 subtask',
     'Orchestrator 对新 subtask 的规划质量依赖 LLM 能力'),
    ('PR 7', 'Token 统计接入',
     'config/llm.py（返回 usage dict）\ntelemetry/execution_log.py（增加 token 字段）',
     'JSONL 日志中 litellm 事件有 prompt_tokens / completion_tokens 字段',
     '无'),
    ('PR 8', '评估是否引入 LangGraph',
     'PR 1-7 完成后，评估：控制流变更是否仍频繁 / 是否需要图结构条件边 / 团队学习成本',
     '团队评估一致决策',
     '引入外部框架有迁移成本，需谨慎'),
]

for pr_name, title, files, acceptance, risk in prs:
    h2(f'{pr_name}：{title}')
    add_table(
        ['项目', '内容'],
        [
            ['修改文件', files],
            ['验收标准', acceptance],
            ['风险', risk],
        ],
        col_widths=[2.5, 14]
    )

# ══════════════════════════════════════════════════════════════
# 18. 待人工确认的问题
# ══════════════════════════════════════════════════════════════
h1('18. 还需要人工确认的问题')
questions = [
    'MAX_PARALLEL_WORKERS=4 在 settings 中存在但未使用，是设计决定还是遗漏？',
    'SubTask.assigned_worker_id 和 SubTask.status 未更新，是有意留给下一阶段，还是遗漏？',
    '_collect_worker_changes 用 copy 而非 merge 覆盖，是否已知这会丢失 Worker 改动？',
    'INTEGRATION_TEST_COMMAND 默认为空，生产环境有计划启用吗？',
    'ENABLE_REVIEW_ROUTER_ACTIVE 默认 False，AgentRouter 实际上从未生效过，是否有启用计划？',
]
for i, q in enumerate(questions, 1):
    bullet(f'{i}. {q}')

# ══════════════════════════════════════════════════════════════
# 附表
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
h1('附表一：当前哪些模块可以直接保留')
add_table(
    ['模块', '是否保留', '原因', '后续改造点'],
    [
        ['agents/claude_code.py', '是', '封装良好，实现 AgentAdapter Protocol', '增加重试逻辑'],
        ['agents/codex.py', '是', '同上，已有 fallback 机制', '无'],
        ['agents/base.py', '是（需小改）', 'AgentTask/AgentResult 抽象合理', '增加 events 字段'],
        ['agents/registry.py', '是', '简单查找表，可扩展', '无'],
        ['agents/context_agent/', '是', '纯 Python，逻辑清晰，无副作用', '无'],
        ['agents/integrator/', '是', '纯规则，无副作用', '无'],
        ['agents/test_agent/', '是', '职责清晰', '无'],
        ['agents/debug_agent/', '是', '职责清晰', '增加失败类型识别'],
        ['agents/worker/workspace.py', '是', '隔离策略设计合理', '无'],
        ['models/ 所有文件', '是（需扩展）', 'Pydantic 数据模型规范', '增加 events、status 写回字段'],
        ['config/settings.py', '是', '配置管理完整', '无'],
        ['config/llm.py', '是（需小改）', '统一 litellm 入口', '增加 token 返回'],
        ['telemetry/execution_log.py', '是', 'JSONL 无损记录', '增加 token 字段'],
    ],
    col_widths=[4.5, 2, 5, 5]
)

h1('附表二：哪些模块必须优先重构')
add_table(
    ['模块', '优先级', '问题', '建议'],
    [
        ['pipeline.py', '1（最高）', '所有控制流硬编码，无法扩展', '提取 StageExecutor，逐步替换为 Orchestrator 决策循环'],
        ['github_agent::_collect_worker_changes', '2', '文件 copy 覆盖导致 Worker 改动丢失', '改用 git merge / 应用 merged_diff'],
        ['agents/orchestrator/', '3', '一次性规划，无反馈循环', '增加 handle_events()，支持追加任务'],
        ['agents/worker/agent.py', '4', '无法上报新发现', 'AgentResult 增加 events 字段，Worker 填写'],
    ],
    col_widths=[4.5, 1.5, 6, 4.5]
)

h1('附表三：从当前架构到目标架构的映射')
add_table(
    ['当前模块', '当前职责', '目标职责', '是否需要拆分'],
    [
        ['pipeline.py', '控制流 + stage 调用 + 重试逻辑', '仅负责启动 + RunState 初始化', '是（拆出 StageExecutor → OrchestratorLoop）'],
        ['OrchestratorAgent', '一次性生成 TaskGraph', '持续决策：接收事件 → 更新计划 → 决定下一步', '否（扩展现有类）'],
        ['ClaudeCodeWorker', '执行 subtask，返回 patch', '执行 subtask，返回 patch + AgentEvents', '否（扩展现有类）'],
        ['ReviewAgent', '审查 patch，返回 approved/block', '审查 patch，返回精确到 subtask 的反馈', '否（扩展现有类）'],
        ['AggregatorAgent', '合并 patch', '同上', '否'],
        ['GitHubAgent._collect_worker_changes', '文件 copy（有 Bug）', 'git merge / apply merged_diff', '否（修复方法）'],
        ['（不存在）', '—', 'RunStateManager（RunState 持久化）', '新增'],
        ['（不存在）', '—', 'AgentEvent（新发现上报协议）', '新增'],
        ['（不存在）', '—', 'TaskBoard（运行时 SubTask 状态管理）', '新增'],
    ],
    col_widths=[3.5, 4, 4.5, 4.5]
)

# ── 保存 ──────────────────────────────────────────────────────
out = r'C:\Users\Administrator\Desktop\OpenClaw架构盘点报告.docx'
doc.save(out)
print(f'文档已生成：{out}')
